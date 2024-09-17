from dotenv import load_dotenv
import random
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import os
import requests

load_dotenv()
token = os.getenv("TOKEN")

api_url = 'https://api.api-ninjas.com/v1/quotes'
response = requests.get(api_url, headers={'X-Api-Key': token})
if response.status_code == requests.codes.ok:
    print("200")
else:
    print("Error:", response.status_code, response.text)

quote_json = response.json()


def set_table_borders(table):
    tbl = table._tbl  # Get the tbl XML element
    tblPr = tbl.tblPr  # Get the table properties
    
    # Create tblBorders element
    tblBorders = OxmlElement('w:tblBorders')
    
    # Define the border attributes
    border_attrs = {
        'top': 'single',
        'left': 'single',
        'bottom': 'single',
        'right': 'single',
        'insideH': 'single',
        'insideV': 'single'
    }

    for border, style in border_attrs.items():
        border_element = OxmlElement(f'w:{border}')
        border_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', style)
        border_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '4')  # Size of the border (1/8 pt)
        border_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')
        tblBorders.append(border_element)

    tblPr.append(tblBorders)



if os.path.exists("patristocrats.docx"):
    os.remove("patristocrats.docx")

doc = Document()

data = quote_json[0]['quote'].lower()
data = "".join([char for char in data if char not in [" ", ",", ".", "'", "-", '"', "+", "-"]])

alphabet = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']

shuffled_alphabet = []

count = 0
added_spaces = 0
num = 0

for i in range(len(data)):
    count += 1
    if count % 5 == 0:
        data = data[:count + added_spaces] + " " + data[count + added_spaces:]
        added_spaces += 1

fails = 0
def shuffle_alphabet(alphabet):
    global shuffled_alphabet
    global fails
    while True:
        shuffled_alphabet = random.sample(alphabet, len(alphabet))
        if all(shuffled_alphabet[i] != alphabet[i] for i in range(len(alphabet))):
            break
        fails += 1
        
shuffle_alphabet(alphabet)

for i in alphabet:
    data = data.replace(i.lower(), shuffled_alphabet[alphabet.index(i)])

num += 1

para = doc.add_paragraph()
run = para.add_run(f"{num}. {data}")
para.paragraph_format.line_spacing_rule = WD_PARAGRAPH_ALIGNMENT.LEFT 
para.paragraph_format.line_spacing = Pt(36)

run.font.name = 'Arial'
run.font.size = Pt(14)

doc.add_paragraph("\n")

table = doc.add_table(rows=3, cols=27)  # +1 for the row labels

hdr_cells = table.rows[0].cells
paragraph = hdr_cells[0].add_paragraph()  # Add a new paragraph to the cell

# Add bold text to the paragraph
run = paragraph.add_run('C')
run.bold = True
for i in range(len(alphabet)):
    hdr_cells[i + 1].text = alphabet[i]

frequency_row = table.rows[1].cells
paragraph = frequency_row[0].add_paragraph()  # Add a new paragraph to the cell

# Add bold text to the paragraph
run = paragraph.add_run('F')
run.bold = True
for i in range (len(alphabet)):
    frequency_row[i + 1].text = str(data.count(alphabet[i]))

replacement_row = table.rows[2].cells
paragraph = replacement_row[0].add_paragraph()  # Add a new paragraph to the cell

# Add bold text to the paragraph
run = paragraph.add_run('R')
run.bold = True
for i in range (len(alphabet)):
    frequency_row[i + 1].text = str(data.count(alphabet[i]))
for i in range (len(alphabet)):
    replacement_row[i + 1].text = ""

set_table_borders(table)

doc.save("patristocrats.docx")

print(fails)